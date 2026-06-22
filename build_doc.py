# -*- coding: utf-8 -*-
import sys, os
sys.path.insert(0, '/Users/chenyutong/.codex/plugins/cache/openai-primary-runtime/documents/26.513.11550/skills/documents/scripts')
from docx import Document
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from table_geometry import apply_table_geometry

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Define styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for i, (name, size, color, before, after) in enumerate([
    ('Heading 1', 16, '2E74B5', 18, 10),
    ('Heading 2', 13, '2E74B5', 14, 7),
    ('Heading 3', 12, '1F4D78', 10, 5),
], 1):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Calibri'
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor.from_string(color)
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(before)
    hs.paragraph_format.space_after = Pt(after)
    hs.paragraph_format.line_spacing = 1.25

# === Helper functions ===
def add_para(text, bold=False, italic=False, style_name='Normal', font_size=None, color=None, alignment=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size: run.font.size = Pt(font_size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    if alignment is not None: p.alignment = alignment
    return p

def make_table(headers, rows, col_weights=None):
    ncols = len(headers)
    weights = col_weights or [1.0] * ncols
    widths = [int(9360 * w / sum(weights)) for w in weights]
    widths[-1] += 9360 - sum(widths)

    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    apply_table_geometry(table, widths)

    # Style
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'B0B0B0')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Header row
    header_fill = 'E8EEF5'
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string('1F4D78')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), header_fill)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return table

# === TITLE ===
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title_p.add_run('AI 算力先进制程——光学互联与先进封装\n材料赛道全景图')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string('1F3A5F')
run.font.name = 'Calibri'

add_para('2026 年 5 月 · 研究参考', italic=True, font_size=10, color='666666')

add_para(
    '核心逻辑：AI 算力从 800G → 1.6T → 3.2T 加速迭代，传统铜互联和有机基板撞上物理极限，'
    '光学互联（光进铜退）和玻璃基封装（玻璃替塑料）两条材料革命同时发生。'
    '先进封装在芯片总成本中的占比已从 2019 年的 15% 飙升到 2025 年的 35-40%。'
    '本文覆盖光学互联、先进封装两大领域涉及的 16 个材料环节及关键标的。',
    font_size=10, color='555555'
)

doc.add_paragraph()  # spacer

# === SECTION 1: 光芯片三大衬底材料 ===
doc.add_heading('一、光芯片三大衬底材料', level=1)
add_para('光模块每升级一代，光芯片成本占比从 30% 飙升到 60-70%，而光芯片的根基是衬底材料。目前三条技术路线并存。')

doc.add_heading('1.1 磷化铟 InP（800G/1.6T 绝对主流）', level=2)
add_para('全球缺口超 70%，日本住友/AXT 垄断 90% 产能，价格一年涨 187%，Lumentum 产能已锁到 2028 年。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['云南锗业', '002428', '国内唯一 6 英寸 InP 衬底量产', '子公司鑫耀半导体扩产至 45 万片/年，国产化率不到 5%，绑定华为/中际旭创'],
        ['三安光电', '600703', '化合物半导体 IDM 龙头', 'GaAs/InP/GaN 全材料体系覆盖，CPO 光引擎激光光源底仓标的'],
        ['有研新材', '600206', '6 英寸 InP 衬底已量产', '多技术路线协同'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

doc.add_heading('1.2 薄膜铌酸锂 TFLN（1.6T→3.2T 未来方案）', level=2)
add_para('电光系数是硅的 100 倍，功耗比硅光低 30%，3.2T 以后可能取代 EML。2026 年是量产拐点元年。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['光库科技', '300620', 'TFLN 调制器器件端最纯正标的', '已具备 800G 相干用 96G/130G TFLN 调制器能力，Q1 扣非 +561%，但估值已透支未来三年'],
        ['天通股份', '600330', 'TFLN 上游晶体和晶片', '国内 4-8 寸铌酸锂晶圆市占率超 50%，但光伏/磁材业务拖累利润表'],
        ['德科立', '688205', '参股铌奥光电切入 TFLN', '弹性最大但基本面弱（2025 净利降 28%），股价一年涨 10 倍'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

doc.add_heading('1.3 硅光 SOI 硅片（800G→1.6T 渗透率飙升）', level=2)
add_para('硅光方案在 1.6T 渗透率已达 30-40%，3.2T 有望提到 50% 以上。SOI 硅片是硅光芯片的基底。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['沪硅产业', '688126', '国内 SOI 硅片最纯正标的', '12 英寸 SOI 硅片通过中际旭创/光迅认证，子公司新傲科技有 Soitec 技术授权'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 2: CPO 共封装光学 ===
doc.add_heading('二、CPO 共封装光学——硅光引擎封装材料', level=1)
add_para('CPO（Co-Packaged Optics）把光引擎直接跟交换芯片封装在同一基板上，省掉两次光电转换，功耗降低 50%，延迟降低 90%。3.2T 时代光互联的终局方案。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['天孚通信', '300394', 'CPO 光引擎器件最纯正', '全球唯二批量供应 CPO 光引擎 FAU（光纤阵列单元），英伟达/Mellanox 核心供应商'],
        ['源杰科技', '688498', 'CW 硅光光源芯片', '国内少数能做 100mW CW 大功率激光器芯片，华为哈勃投资，已通过中际旭创验证'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 3: HBM 专用封装材料 ===
doc.add_heading('三、HBM 专用封装材料', level=1)
add_para('HBM 市场 2024 年 150 亿美元 → 2027 年 576 亿美元，三年 4 倍。每颗 HBM 消耗的先进封装材料价值量约 $30-35。HBM4 引入混合键合（Hybrid Bonding），材料消耗结构将剧变。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['联瑞新材', '688300', '硅微粉（EMC 填料）', '半导体 EMC 用球形硅微粉全球份额 25%，low-α 球硅打破日本电气化学垄断。HBM 堆叠层数增加，EMC 用量线性增加'],
        ['飞凯材料', '300398', '临时键合胶+EMC', 'HBM 临时键合胶全球 >30% 份额。每片晶圆需 2-3 次临时键合/解键合，耗材用量跟堆叠层数成线性关系'],
        ['华海诚科', '688535', 'HBM 用 GMC 颗粒塑封料', '国内唯一通过客户验证的 HBM 专用 GMC（颗粒状环氧塑封料）企业'],
        ['德邦科技', '688035', '芯片级底填胶+固晶胶', 'HBM 堆叠芯片间 Underfill 底填胶，直接跟日本 Namics/日立化成正面竞争'],
        ['雅克科技', '002409', '前驱体+SOD', 'HBM 逻辑芯片 High-K 金属栅极前驱体材料，SK 海力士核心供应商'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 4: HBM4 混合键合材料 ===
doc.add_heading('四、HBM4 混合键合材料', level=1)
add_para('HBM4 堆叠 16-20 层，互连间距要压到 10μm 以下，微凸块物理上做不到了，必须换成混合键合——铜对铜直接连接。全球能提供全套化学品的公司一只手数得过来。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['上海新阳', '300236', 'CMP 后清洗液+电镀液', 'CMP 后清洗液是国内唯一能满足混合键合表面洁净度要求的（Cu 粗糙度 < 0.2nm RMS），绑定长电/通富/长鑫'],
        ['安集科技', '688019', 'CMP 抛光液+清洗', '铜 CMP 抛光液 14nm 以下已导入中芯国际，混合键合用超精密铜抛光液正在验证'],
        ['艾森股份', '688720', 'D2D 互连电镀化学品', '芯片级 D2D 互连微型铜柱电镀，线宽/间距逼近 1μm/1μm'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 5: ABF 载板及上游材料 ===
doc.add_heading('五、ABF 载板及上游材料', level=1)
add_para('CoWoS 封装瓶颈卡在 ABF 载板——每颗 B200 GPU 消耗一片大尺寸 ABF 载板，2025-2027 年缺口约 20-30 万片/月。ABF 膜被日本味之素独家垄断，国产替代机会集中在载板制造、上游树脂和超薄玻纤布。')
doc.add_heading('5.1 ABF 载板制造', level=2)
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['深南电路', '002916', 'FC-BGA 载板国内最领先', '14 层及以下 FC-BGA 载板 2025Q1 开始放量，全年预计 5-6 亿营收，质量对标日本揖斐电'],
        ['兴森科技', '002436', 'BT/ABF 载板双线', '广州 ABF 载板一期 2025 年产能逐月爬坡，珠海 BT 载板满产，三星/海思认证中'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

doc.add_heading('5.2 ABF 膜上游——特种树脂 + 低介电玻纤布', level=2)
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['圣泉集团', '605589', 'BT 氰酸酯树脂', '国内唯一实现 BT 树脂量产，PPO 树脂已批量供货生益科技（年收入 6-7 亿），ABF 膜级树脂送样测试中'],
        ['东材科技', '601208', '电子级马来酰亚胺树脂', '打破日韩垄断，电子级特种环氧＋马来酰亚胺树脂双线，光学膜＋质子交换膜＋电子树脂三引擎'],
        ['宏昌电子', '603002', '电子级环氧树脂', '高频高速覆铜板用环氧树脂国内领先，ABF 膜级环氧树脂正在研发'],
        ['宏和科技', '603256', '超薄电子布（4μm）', '全球少数能做 12μm 以下超薄电子布的厂商，4μm 打破日本旭化成垄断，ABF 载板补强材料送样中'],
        ['中国巨石', '600176', '低介电 Low-Dk 玻纤', '电子级细纱全球份额 25%，Low-Dk 低介电玻纤已批量供货生益科技/台光'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 6: TSV/RDL 互连材料 ===
doc.add_heading('六、TSV / RDL 互连材料', level=1)
add_para('先进封装从 2D 走向 2.5D/3D，硅中介层上的 RDL（再分布层）和 TSV（硅通孔）需要大量高纯度电镀液、低介电绝缘介质和微球焊料。')

doc.add_heading('6.1 电镀液 + 清洗液', level=2)
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['上海新阳', '300236', '电镀液+清洗液绝对龙头', '硫酸铜电镀液突破台积电 5nm 验证（国内唯一），用于 TSV/RDL 铜互连和 Bumping 铜柱'],
        ['艾森股份', '688720', '先进封装电镀液全品类', '电镀铜基液、锡银电镀液、镍钯金电镀液全覆盖，卡位 PSPI 光刻胶和 HBM 供应链'],
        ['天承科技', '688603', 'PCB/载板水平沉铜', 'PCB 水平沉铜化学品国内第一，正向 FC-BGA 载板电镀切入'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

doc.add_heading('6.2 RDL 低介电绝缘介质', level=2)
add_para('RDL 线宽/线距从 10μm/10μm 向 2μm/2μm 推进，绝缘介质介电常数越低越好。目前 PID 光敏聚酰亚胺、PBO 聚苯并噁唑、BCB 苯并环丁烯三条路线几乎 100% 进口。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['鼎龙股份', '300054', 'PSPI 光敏聚酰亚胺', '显示用 PSPI 已规模出货，半导体封装用 PSPI 正在晶圆厂验证，CMP 抛光垫+抛光液+PI 三线并进'],
        ['八亿时空', '688181', 'PI 单体+液晶', 'PI 单体已稳定供货全球头部 PI 厂商，向封装级 PI 薄膜/浆料延伸'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

doc.add_heading('6.3 微电子焊粉', level=2)
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['有研粉材', '688456', '微电子焊粉', '锡基焊粉国内龙头，BGA/CSP 封装用微球焊粉粒径可做到 5-25μm，先进封装 bumping 用超细焊粉国产替代稀缺标的'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 7: 芯片粘接 / 烧结材料 ===
doc.add_heading('七、芯片粘接 / 烧结材料（Die Attach）', level=1)
add_para('AI 芯片裸 die 面积 800mm²+，热流密度 150W/cm²，传统银胶导热率不够。银烧结（Ag Sintering）纳米银膏在 250°C 下烧结，导热率是传统银胶的 10 倍，SiC 功率模块和 AI GPU 都在转。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['德邦科技', '688035', '固晶胶+银烧结国内第一', '纳米银烧结膏已用于华为昇腾/海光 AI 芯片封装，2025 年银烧结收入增速 200%+。TIM1+Underfill+Die Attach+银烧结四品线'],
        ['赛伍技术', '603212', 'IGBT/SiC 烧结银膏', '光伏背板龙头转型，烧结银膏已批量出货汇川/中车，AI 服务器芯片烧结膏送样中'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 8: 光模块导热 / 热界面材料 ===
doc.add_heading('八、光模块导热 / 热界面材料（TIM）', level=1)
add_para('光模块速率每翻一倍，单模块功耗增加 50-60%，1.6T 光模块单颗功耗 ~30W，传统导热硅脂不够用。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['德邦科技', '688035', '芯片级 TIM1/TIM2', '导热凝胶/导热垫片/导热硅脂全覆盖，TIM1 导热率 8-12W/m·K 国内最高，客户覆盖华为海思/新华三/光迅'],
        ['中石科技', '300684', '人工合成石墨散热膜', '高导热石墨膜国内出货量最大，VC 均温板/热管全品类，TIM 导热胶已切入英伟达 GB300 验证'],
        ['飞荣达', '300602', 'EMI 屏蔽+导热双龙头', '英伟达 GB300 液态金属导热方案核心供应商（单机柜 1.5-2 万元），华为昇腾液冷屏蔽一体化方案独家'],
        ['回天新材', '300041', '导热结构胶', '5G 基站导热胶国产份额第一，光模块/交换机导热材料正在送样认证'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 9: 先进封测特种气体 ===
doc.add_heading('九、先进封测特种气体', level=1)
add_para('先进封装（尤其是 TSV 深硅刻蚀和 CVD 介质沉积）每片晶圆消耗的特种气体价值量是传统封装的 5-8 倍。TSV 刻蚀深度 50-100μm，需要大量高纯度氟基气体。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['华特气体', '688268', '光刻气+刻蚀气全品类', '国内高纯度氟碳类气体（CF₄/CHF₃/C₄F₆/C₄F₈）最全供应商，TSV 深硅刻蚀核心耗材，光刻气通过 ASML 认证'],
        ['昊华科技', '600378', '含氟电子特气全品类', '中化集团旗下，NF₃（清洗气）、SF₆（刻蚀气）、CF₄ 全品类，TSV 刻蚀气体供应商'],
        ['凯美特气', '002549', '超高纯电子特气', 'CO₂/CO/H₂/CH₄ 超高纯特气，先进封装 CVD 工艺气源'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 10: 玻璃基封装 TGV ===
doc.add_heading('十、玻璃基封装 TGV', level=1)
add_para('传统有机载板在 2.5D/3D 封装中遇到热膨胀系数不匹配、翘曲、高密度布线瓶颈。玻璃基封装（Glass Core Substrate）通孔密度提升 10 倍、热稳定性提升一个数量级。英特尔 2030 年目标商用。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['沃格光电', '603773', '国内 TGV 技术最领先', 'TGV 玻璃通孔技术国内首发，PVD 镀铜+湿法蚀刻自研，与华为/中芯国际联合研发，已规划年产 100 万片产能'],
        ['凯盛科技', '600552', '玻璃基板原片', '显示玻璃基板国内第三，UTG 超薄玻璃折叠屏华为独供，有玻璃原片成本优势'],
        ['大族激光', '002008', 'TGV 激光打孔设备', '超快紫外激光打孔深度/锥度控制国内最强，已为沃格/京东方供 TGV 激光设备'],
        ['帝尔激光', '300776', '光伏激光+TGV 储备', 'TGV 激光诱导深度刻蚀技术有储备，主业仍在光伏'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 11: 功率芯片封装材料 ===
doc.add_heading('十一、功率芯片封装材料', level=1)
add_para('AI 服务器单机柜功耗从 40kW 飙升到 120kW，传统 48V 供电正向 800V HVDC 切换，SiC 功率模块封装和大电流基板需求暴增。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['博敏电子', '603936', 'AMB 陶瓷基板', '国内少数能量产 AMB（活性金属钎焊）陶瓷基板的公司，SiC 功率模块封装核心材料，AMB 出货量全国前三'],
        ['天岳先进', '688234', '8 英寸 SiC 衬底', '全球第三家量产 8 英寸碳化硅衬底的企业（次于 Wolfspeed/Coherent），2025 年营收增 48%'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 12: 散热配套 ===
doc.add_heading('十二、系统级散热配套', level=1)
add_para('AI 单机柜 120kW 已是现实，英伟达 GB300 NVL72 整机柜功率直奔 140kW。风冷极限 ~30kW/柜，其余必须走液冷。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['英维克', '002837', '数据中心液冷系统', '液冷温控绝对龙头（腾讯/字节/阿里/华为主力供应商），液冷渗透率每升 10%，导热材料需求一个量级'],
        ['飞荣达', '300602', '液冷屏蔽一体化', '英伟达 GB300 液态金属导热方案核心供应商'],
        ['中石科技', '300684', '石墨散热+VC 均温板', 'AI 服务器散热全场景覆盖'],
        ['思泉新材', '301489', '人工合成石墨散热膜', '小米/OPPO 核心供应商，正向 AI 服务器导热材料拓展'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === SECTION 13: 洁净室耗材 ===
doc.add_heading('十三、洁净室耗材', level=1)
add_para('先进封装对环境洁净度的要求是 ISO Class 1（每立方米不超过 10 个 0.1μm 颗粒），TSV 刻蚀和 PVD 溅射环节需要大量一次性高洁净耗材。被忽视的品类，但毛利极高（60-80%）。')
make_table(
    ['标的', '代码', '卡位', '核心亮点'],
    [
        ['美埃科技', '688376', '超高效空气过滤', '半导体洁净室 U15/U16 超高效过滤器国内第一，中芯国际/长电/华虹核心供应商，过滤器每 1-2 年更换一次，耗材粘性极高'],
    ],
    col_weights=[0.9, 0.7, 1.4, 3.0]
)

# === 全景速查表 ===
doc.add_heading('十四、全景速查表', level=1)
add_para('以下是本文覆盖的全部 16 个材料环节及关键标的速查索引。')

summary_headers = ['环节', '核心材料', '国产化率', '最纯正标的', '确定性']
summary_rows = [
    ['光芯片衬底', 'InP 磷化铟', '< 5%', '云南锗业', '高（缺口刚需）'],
    ['光芯片衬底', 'TFLN 铌酸锂', '~10%', '天通股份', '中（3.2T 时代放量）'],
    ['光引擎光源', 'CW 硅光大功率激光器', '~15%', '源杰科技', '高'],
    ['HBM 封装', 'EMC 硅微粉填料', '~25%', '联瑞新材', '极高（HBM 层数线性驱动）'],
    ['HBM 封装', '临时键合胶', '~30%', '飞凯材料', '极高'],
    ['HBM 封装', 'GMC 颗粒塑封料', '~5%', '华海诚科', '高'],
    ['HBM4 混合键合', 'CMP 清洗液+电镀液', '~10%', '上海新阳', '中（HBM4 节奏待定）'],
    ['CoWoS 封装', 'ABF 载板', '~30%', '深南电路', '极高'],
    ['ABF 载板上游', 'BT 氰酸酯树脂', '~5%', '圣泉集团', '高'],
    ['ABF 载板上游', '超薄电子布 4μm', '~10%', '宏和科技', '中'],
    ['RDL 绝缘介质', 'PID 光敏聚酰亚胺', '~3%', '鼎龙股份（布局中）', '低（国内几乎空白）'],
    ['TSV/RDL 电镀', '硫酸铜电镀液', '~20%', '上海新阳', '极高'],
    ['TSV 刻蚀', 'C₄F₈/CF₄ 含氟特气', '~30%', '华特气体', '高'],
    ['芯片粘接', '纳米银烧结膏', '~20%', '德邦科技', '高'],
    ['TIM 导热', '导热凝胶/垫片', '~40%', '德邦科技/中石科技', '高'],
    ['玻璃基封装', 'TGV 玻璃通孔', '~5%', '沃格光电', '低（2030 年前慢变量）'],
    ['SiC 功率封装', 'AMB 陶瓷基板', '~25%', '博敏电子', '中'],
    ['系统液冷', '冷板+manifold', '—', '英维克（锚定）', '高'],
    ['洁净环境', '超高效过滤器', '~40%', '美埃科技', '高'],
]
make_table(summary_headers, summary_rows, col_weights=[0.9, 1.4, 0.6, 1.3, 0.8])

# === SAVE ===
output_path = '/Users/chenyutong/Downloads/a-stock-data-main/AI算力先进制程_材料赛道全景图.docx'
doc.save(output_path)
print(f'DOCX saved to: {output_path}')
